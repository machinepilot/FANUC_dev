#!/usr/bin/env python3
"""
Extract text from ESA and BLM .docx manuals for LDJ integration reference.
Requires: pip install python-docx

Usage: python extract_docx.py
Output: archive/Esa_Manual_extracted.txt, archive/BLM_MANUAL_extracted.txt
"""

import re
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

SCRIPT_DIR = Path(__file__).resolve().parent
ARCHIVE_DIR = SCRIPT_DIR / "archive"
ARCHIVE_DIR.mkdir(exist_ok=True)

SOURCES = [
    ("Esa_Manual_P0122.docx", "Esa_Manual_extracted.txt"),
    ("BLM_MANUAL.docx", "BLM_MANUAL_extracted.txt"),
]


def extract_docx_via_zip(docx_path: Path) -> str:
    """Fallback: extract text from document.xml via zip (handles large embedded content)."""
    lines = []
    with zipfile.ZipFile(docx_path, "r") as z:
        with z.open("word/document.xml") as f:
            content = f.read().decode("utf-8", errors="replace")
    # Extract text from <w:t> tags, split by paragraphs
    for para in re.split(r"<w:p\b", content):
        texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", para)
        if texts:
            line = "".join(texts).strip()
            if line:
                lines.append(line)
    return "\n\n".join(lines)


def get_heading_level(paragraph):
    """Return heading level (1-9) or 0 if not a heading."""
    pPr = paragraph._element.get_or_add_pPr()
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return 0
    val = pStyle.get(qn("w:val"))
    if val and val.startswith("Heading"):
        try:
            return int(val.replace("Heading", "").strip() or "1")
        except ValueError:
            return 1
    return 0


def extract_document(docx_path: Path) -> str:
    """Extract text from .docx preserving structure (headings, paragraphs, tables)."""
    doc = Document(docx_path)
    lines = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = element
            p = next((p for p in doc.paragraphs if p._element == para), None)
            if p is None:
                continue
            level = get_heading_level(p)
            text = p.text.strip()
            if not text:
                continue
            if level > 0:
                prefix = "#" * level + " "
                lines.append(f"\n{prefix}{text}\n")
            else:
                lines.append(text)

        elif tag == "tbl":
            # Extract table
            table = next((t for t in doc.tables if t._element == element), None)
            if table is not None:
                lines.append("")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                    if row_text.strip():
                        lines.append("| " + row_text + " |")
                lines.append("")

    return "\n".join(lines).strip()


def main():
    for src_name, out_name in SOURCES:
        src_path = SCRIPT_DIR / src_name
        out_path = ARCHIVE_DIR / out_name

        if not src_path.exists():
            print(f"Skip: {src_name} not found")
            continue

        print(f"Extracting {src_name} -> {out_name}")
        try:
            text = extract_document(src_path)
            out_path.write_text(text, encoding="utf-8")
            print(f"  Wrote {len(text)} chars to {out_path}")
        except Exception as e:
            print(f"  python-docx failed ({e}), trying zip fallback...")
            try:
                text = extract_docx_via_zip(src_path)
                out_path.write_text(text, encoding="utf-8")
                print(f"  Wrote {len(text)} chars to {out_path} (via zip)")
            except Exception as e2:
                print(f"  Error: {e2}")
                sys.exit(1)

    print("Done.")


if __name__ == "__main__":
    main()
